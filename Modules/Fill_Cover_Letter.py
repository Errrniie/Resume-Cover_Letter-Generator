#!/usr/bin/env python3
"""
Fill a cover letter .docx template from JSON and save to a caller-provided path.

Flat text replacements only (no bullets or skills). Does not choose paths via Main.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from dataclasses import dataclass, field
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document
from docx.text.paragraph import Paragraph

from datetime import date

from config.Cover_Letter_Config import PLACEHOLDER_ALIASES, REQUIRED_COVER_LETTER_FIELDS
from config.Cover_Letter_Validator import validate_cover_letter_data
from config.Loader import load_config, resolve_template_path

PLACEHOLDER_RE = re.compile(r"^\{\{\s*(.+?)\s*\}\}$")
COVER_LETTER_FILENAME_PREFIX = "Ernesto_Carlton_Cover_Letter"
INVALID_FILENAME_CHARS = re.compile(r'[<>:"/\\|?*\x00-\x1f]')


@dataclass
class CoverLetterData:
    """Normalized cover letter field values for template replacement."""

    fields: dict[str, str] = field(default_factory=dict)


def blank_value(value: object) -> bool:
    return value is None or (isinstance(value, str) and not value.strip())


def as_text(value: object) -> str:
    if blank_value(value):
        return ""
    return str(value).strip()


def placeholder_token(key: str) -> str:
    match = PLACEHOLDER_RE.match(key.strip())
    inner = match.group(1) if match else key.strip()
    return f"{{{{ {inner} }}}}"


def sanitize_filename_part(text: str) -> str:
    cleaned = INVALID_FILENAME_CHARS.sub("", text)
    cleaned = re.sub(r"\s+", "_", cleaned.strip())
    cleaned = re.sub(r"_+", "_", cleaned)
    return cleaned.strip("._")


def format_letter_date(run_date: date | None = None) -> str:
    """Today's date as e.g. May 21, 2026 (no leading zero on day)."""
    today = run_date or date.today()
    return f"{today.strftime('%B')} {today.day}, {today.year}"


def suggested_cover_letter_filename(raw: dict) -> str:
    """Ernesto_Carlton_Cover_Letter_{company}_{position}.docx (for CLI/docs)."""
    company = sanitize_filename_part(str(raw.get("company", "")))
    position = sanitize_filename_part(str(raw.get("position", "")))
    if not company or not position:
        raise ValueError(
            'JSON must include non-empty "company" and "position" for output naming'
        )
    return f"{COVER_LETTER_FILENAME_PREFIX}_{company}_{position}.docx"


def parse_cover_letter_data(raw: dict) -> CoverLetterData:
    """Normalize cover letter fields from a parsed JSON object."""
    if not isinstance(raw, dict):
        raise ValueError(f"JSON root must be an object, got {type(raw).__name__}")

    fields: dict[str, str] = {}

    for key, value in raw.items():
        if key not in REQUIRED_COVER_LETTER_FIELDS:
            continue
        if blank_value(value):
            fields[key] = ""
        elif not isinstance(value, (str, int, float, bool)):
            raise ValueError(f'"{key}" must be a string or scalar value')
        else:
            fields[key] = as_text(value)

    return CoverLetterData(fields=fields)


def load_json(json_path: Path) -> CoverLetterData:
    with json_path.open(encoding="utf-8") as f:
        raw = json.load(f)
    return parse_cover_letter_data(raw)


def build_replacements(
    data: CoverLetterData,
    *,
    letter_date: str | None = None,
) -> dict[str, str]:
    # position is required in JSON for folder/filename alignment; templates may omit {{ position }}.
    replacements: dict[str, str] = {}
    for key, value in data.fields.items():
        replacements[placeholder_token(key)] = value
        for alias in PLACEHOLDER_ALIASES.get(key, ()):
            replacements[alias] = value
    replacements[placeholder_token("date")] = letter_date or format_letter_date()
    return replacements


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(text)
        return
    for run in runs[1:]:
        run._element.getparent().remove(run._element)
    runs[0].text = text


def replace_in_paragraph(paragraph: Paragraph, replacements: dict[str, str]) -> None:
    text = paragraph.text
    if not text or "{{" not in text:
        return

    updated = text
    for token, value in replacements.items():
        if token in updated:
            updated = updated.replace(token, value)

    if updated == text:
        return

    updated = updated.rstrip("\t\n\r ")
    set_paragraph_text(paragraph, updated)


def iter_all_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def apply_replacements(doc: Document, replacements: dict[str, str]) -> None:
    for paragraph in iter_all_paragraphs(doc):
        replace_in_paragraph(paragraph, replacements)


def fill_cover_letter(
    template_path: str | Path,
    output_path: str | Path,
    json_path: str | Path | None = None,
    *,
    data: CoverLetterData | None = None,
) -> Path:
    """
    Copy template, replace {{ field }} placeholders, and save to output_path.

    Caller supplies template_path and output_path. Does not use Application_Path or Main.
    """
    template_path = Path(template_path).resolve()
    output_path = Path(output_path).resolve()

    if data is None:
        if json_path is None:
            raise ValueError("fill_cover_letter requires json_path or data")
        json_path = Path(json_path).resolve()
        if not json_path.is_file():
            raise FileNotFoundError(f"JSON file not found: {json_path}")
        data = load_json(json_path)

    if not template_path.is_file():
        raise FileNotFoundError(f"Template not found: {template_path}")
    if output_path.suffix.lower() != ".docx":
        raise ValueError(f"Output path must be a .docx file: {output_path}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output_path)

    doc = Document(str(output_path))
    apply_replacements(doc, build_replacements(data))
    doc.save(str(output_path))

    return output_path


def _default_template_path() -> Path:
    cfg = load_config()
    template = cfg.template("cover_letter")
    if template is None or not template.enabled or template.path is None:
        raise ValueError(
            "Cover letter template is not enabled or has no path in settings.json"
        )
    return resolve_template_path(template)


def _parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Fill a cover letter .docx from JSON (standalone CLI).",
    )
    parser.add_argument("json", type=Path, help="Path to cover letter JSON data.")
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        required=True,
        help="Output .docx path.",
    )
    parser.add_argument(
        "--template",
        type=Path,
        default=None,
        help="Cover letter .docx template (default: settings.json templates.cover_letter).",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = _parse_args(argv)

    try:
        with args.json.open(encoding="utf-8") as f:
            raw = json.load(f)
    except (OSError, json.JSONDecodeError) as exc:
        print(f"Error loading JSON: {exc}", file=sys.stderr)
        return 1

    result = validate_cover_letter_data(raw)
    if not result.valid:
        print(f"Validation failed: {args.json}", file=sys.stderr)
        for message in result.messages():
            print(f"  - {message}", file=sys.stderr)
        return 1

    try:
        template_path = args.template or _default_template_path()
        out = fill_cover_letter(template_path, args.output, data=parse_cover_letter_data(raw))
    except (FileNotFoundError, ValueError, RuntimeError) as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1

    print(f"Created: {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
