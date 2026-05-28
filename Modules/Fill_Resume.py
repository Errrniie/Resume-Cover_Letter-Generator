#!/usr/bin/env python3
"""
Fill a resume .docx template from JSON and save to a caller-provided path.

Does not choose template, output location, or filename — pass all paths in from Main.
"""

from __future__ import annotations

import json
import re
import shutil
from collections import defaultdict
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

PLACEHOLDER_RE = re.compile(r"^\{\{\s*(.+?)\s*\}\}$")
BULLET_IN_TEXT_RE = re.compile(r"\{\{\s*(\w+)_bullet_(\d+)\s*\}\}")
FLAT_BULLET_KEY_RE = re.compile(r"^(\w+)_bullet_(\d+)$")
MAX_SKILL_CATEGORIES = 3
MAX_SKILLS_PER_CATEGORY = 6


@dataclass
class ResumeData:
    """Normalized values extracted from a resume JSON file."""

    replacements: dict[str, str] = field(default_factory=dict)
    bullet_groups: dict[str, list[str]] = field(default_factory=dict)
    skill_categories: list[dict[str, list[str] | str]] = field(default_factory=list)


def blank_value(value: object) -> bool:
    return value is None or (isinstance(value, str) and not value.strip())


def as_text(value: object) -> str:
    if blank_value(value):
        return ""
    return str(value).strip()


def placeholder_token(key: str) -> str:
    match = PLACEHOLDER_RE.match(key.strip())
    inner = match.group(1) if match else key.strip()
    return f"{{{{ {inner} }}}}"


def parse_resume_data(raw: dict) -> ResumeData:
    """Normalize resume data from a parsed JSON object."""
    if not isinstance(raw, dict):
        raise ValueError(f"JSON root must be an object, got {type(raw).__name__}")

    data = ResumeData()
    flat: dict[str, object] = {}
    flat_bullets: dict[str, dict[int, str]] = defaultdict(dict)

    for key, value in raw.items():
        if key == "skills":
            if not isinstance(value, list):
                raise ValueError('"skills" must be an array of category objects')
            data.skill_categories = _parse_skill_categories(value)
            continue
        if key.endswith("_bullets"):
            if not isinstance(value, list):
                raise ValueError(f'"{key}" must be an array of strings')
            prefix = key[: -len("_bullets")]
            data.bullet_groups[prefix] = [as_text(item) for item in value if not blank_value(item)]
            continue
        flat[key] = value

    for key, value in flat.items():
        match = FLAT_BULLET_KEY_RE.match(key)
        if match:
            prefix, index = match.group(1), int(match.group(2))
            text = as_text(value)
            if text:
                flat_bullets[prefix][index] = text
            continue
        if blank_value(value):
            data.replacements[key] = ""
        else:
            data.replacements[key] = as_text(value)

    for prefix, indices in flat_bullets.items():
        if prefix not in data.bullet_groups:
            data.bullet_groups[prefix] = [indices[i] for i in sorted(indices)]

    data.replacements.update(_skill_replacements(data.skill_categories))
    return data


def load_json(json_path: Path) -> ResumeData:
    """Read and normalize resume data from a JSON file."""
    with json_path.open(encoding="utf-8") as f:
        raw = json.load(f)
    return parse_resume_data(raw)


def _parse_skill_categories(value: list) -> list[dict[str, list[str] | str]]:
    categories: list[dict[str, list[str] | str]] = []
    for item in value:
        if not isinstance(item, dict):
            raise ValueError('Each skills entry must be an object with "category" and "skills"')
        category = as_text(item.get("category", item.get("name", "")))
        skills_raw = item.get("skills", [])
        if not isinstance(skills_raw, list):
            raise ValueError('"skills" must be an array of strings')
        skills = [as_text(skill) for skill in skills_raw if not blank_value(skill)]
        if category or skills:
            categories.append({"category": category, "skills": skills})
    return categories[:MAX_SKILL_CATEGORIES]


def _skill_replacements(categories: list[dict[str, list[str] | str]]) -> dict[str, str]:
    replacements: dict[str, str] = {}
    for row_idx, entry in enumerate(categories, start=1):
        replacements[f"skill_cat_{row_idx}"] = str(entry["category"])
        skills = entry["skills"]
        if not isinstance(skills, list):
            continue
        for col_idx, skill in enumerate(skills[:MAX_SKILLS_PER_CATEGORY], start=1):
            replacements[f"skill_{row_idx}{col_idx}"] = str(skill)
        for col_idx in range(len(skills) + 1, MAX_SKILLS_PER_CATEGORY + 1):
            replacements[f"skill_{row_idx}{col_idx}"] = ""
    return replacements


def build_replacements(data: ResumeData) -> dict[str, str]:
    replacements: dict[str, str] = {}
    for key, value in data.replacements.items():
        replacements[placeholder_token(key)] = value
    for prefix, bullets in data.bullet_groups.items():
        for index, text in enumerate(bullets, start=1):
            replacements[placeholder_token(f"{prefix}_bullet_{index}")] = text
    return replacements


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(text)
        return
    for run in runs[1:]:
        run._element.getparent().remove(run._element)
    runs[0].text = text


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_bullet_paragraph_after(template: Paragraph, text: str) -> Paragraph:
    """Insert a new bullet paragraph after template, copying its formatting."""
    new_element = deepcopy(template._element)
    template._element.addnext(new_element)
    new_paragraph = Paragraph(new_element, template._parent)
    set_paragraph_text(new_paragraph, text.rstrip("\t\n\r "))
    return new_paragraph


def replace_in_paragraph(paragraph: Paragraph, replacements: dict[str, str]) -> None:
    text = paragraph.text
    if not text or "{{" not in text:
        return

    updated = text
    for token, value in replacements.items():
        if token in updated:
            updated = updated.replace(token, value)

    if updated == text:
        return

    updated = updated.rstrip("\t\n\r ")
    set_paragraph_text(paragraph, updated)


def iter_body_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph


def iter_all_paragraphs(doc: Document):
    for paragraph in iter_body_paragraphs(doc):
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def find_skills_table(doc: Document) -> Table | None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "skill_cat" in cell.text:
                    return table
    return None


def delete_table_row(table: Table, row_index: int) -> None:
    row = table.rows[row_index]
    table._tbl.remove(row._tr)


def collect_bullet_paragraphs(doc: Document) -> dict[str, list[tuple[Paragraph, str]]]:
    groups: dict[str, list[tuple[Paragraph, str]]] = defaultdict(list)
    for paragraph in iter_body_paragraphs(doc):
        matches = BULLET_IN_TEXT_RE.findall(paragraph.text)
        if not matches:
            continue
        prefix, number = matches[0]
        token = placeholder_token(f"{prefix}_bullet_{number}")
        groups[prefix].append((paragraph, token))
    return groups


def process_bullets(doc: Document, bullet_groups: dict[str, list[str]]) -> None:
    grouped = collect_bullet_paragraphs(doc)
    to_delete: list[Paragraph] = []

    for prefix, paragraphs in grouped.items():
        bullets = bullet_groups.get(prefix, [])
        template_count = len(paragraphs)

        for index in range(min(template_count, len(bullets))):
            paragraph, token = paragraphs[index]
            updated = paragraph.text
            if token in updated:
                updated = updated.replace(token, bullets[index])
                updated = updated.rstrip("\t\n\r ")
                set_paragraph_text(paragraph, updated)

        for index in range(len(bullets), template_count):
            to_delete.append(paragraphs[index][0])

        if len(bullets) > template_count and template_count > 0:
            anchor = paragraphs[template_count - 1][0]
            for extra_index in range(template_count, len(bullets)):
                anchor = insert_bullet_paragraph_after(anchor, bullets[extra_index])

    for paragraph in reversed(to_delete):
        delete_paragraph(paragraph)


def process_skills_table(doc: Document, categories: list[dict[str, list[str] | str]]) -> None:
    table = find_skills_table(doc)
    if table is None:
        return

    for row_index in range(len(table.rows) - 1, len(categories) - 1, -1):
        delete_table_row(table, row_index)


def apply_replacements(doc: Document, replacements: dict[str, str]) -> None:
    for paragraph in iter_all_paragraphs(doc):
        replace_in_paragraph(paragraph, replacements)


def fill_resume(
    template_path: str | Path,
    output_path: str | Path,
    json_path: str | Path | None = None,
    *,
    data: ResumeData | None = None,
) -> Path:
    """
    Copy template, fill placeholders from JSON, and save to output_path.

    Args:
        template_path: Path to the .docx template (not modified).
        output_path: Full path for the filled .docx to write.
        json_path: Path to the resume JSON data file (omit if data is provided).
        data: Pre-parsed resume data (e.g. from another tool via Main).

    Returns:
        Resolved path to the saved document.
    """
    template_path = Path(template_path).resolve()
    output_path = Path(output_path).resolve()

    if data is None:
        if json_path is None:
            raise ValueError("fill_resume requires json_path or data")
        json_path = Path(json_path).resolve()
        if not json_path.is_file():
            raise FileNotFoundError(f"JSON file not found: {json_path}")
        data = load_json(json_path)
    elif json_path is not None:
        json_path = Path(json_path).resolve()

    if not template_path.is_file():
        raise FileNotFoundError(f"Template not found: {template_path}")
    if output_path.suffix.lower() != ".docx":
        raise ValueError(f"Output path must be a .docx file: {output_path}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output_path)

    doc = Document(str(output_path))
    process_skills_table(doc, data.skill_categories)
    process_bullets(doc, data.bullet_groups)
    apply_replacements(doc, build_replacements(data))
    doc.save(str(output_path))

    return output_path
